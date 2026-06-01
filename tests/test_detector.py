"""
Tests for FormatDetector — format auto-detection and data extraction.
Covers: CSV, JSON, DOCX, XLSX, XML, TXT formats.
"""

import json
import os

import pytest

from src.detector import FormatDetector
from src.models import FileFormat

# ===========================================================================
# Format Detection Tests
# ===========================================================================


class TestFormatDetection:
    """Test that FormatDetector correctly identifies file formats from extensions."""

    @pytest.mark.parametrize(
        "filename,expected",
        [
            ("data.csv", FileFormat.CSV),
            ("data.json", FileFormat.JSON),
            ("data.docx", FileFormat.DOCX),
            ("data.xlsx", FileFormat.XLSX),
            ("data.xml", FileFormat.XML),
            ("data.txt", FileFormat.TXT),
            ("DATA.CSV", FileFormat.CSV),  # uppercase extension
            ("Data.Json", FileFormat.JSON),  # mixed case
        ],
    )
    def test_detect_format_by_extension(self, filename, expected):
        fmt = FormatDetector.detect_format(filename)
        assert fmt == expected

    def test_detect_format_unknown_defaults_to_txt(self):
        """Unknown extensions should default to TXT."""
        assert FormatDetector.detect_format("data.xyz") == FileFormat.TXT
        assert FormatDetector.detect_format("data") == FileFormat.TXT
        assert FormatDetector.detect_format("data.parquet") == FileFormat.TXT

    def test_mime_map_completeness(self):
        """MIME_MAP should cover all FileFormat enum values."""
        detected_mimes = set(FormatDetector.MIME_MAP.values())
        for fmt in FileFormat:
            assert fmt in detected_mimes, f"FileFormat.{fmt} not in MIME_MAP"


# ===========================================================================
# CSV Extraction Tests
# ===========================================================================


class TestCSVExtraction:
    """Test CSV file extraction."""

    def test_extract_csv_basic(self, csv_file, sample_records):
        records = FormatDetector.extract(csv_file)
        assert len(records) == len(sample_records)
        assert records[0]["full_name"] == "Muhammad Tayyab"
        assert records[1]["email"] == "ali.ahmed@test.org"

    def test_extract_csv_strips_whitespace(self, sample_dir):
        """CSV extractor should strip whitespace from keys and values."""
        path = os.path.join(sample_dir, "spaces.csv")
        with open(path, "w", newline="") as f:
            f.write(" name , dob , email \n")
            f.write("  Test User , 1990-01-01 , test@test.com \n")
        records = FormatDetector.extract(path)
        assert len(records) == 1
        assert "name" in records[0]
        assert records[0]["name"] == "Test User"
        assert records[0]["email"] == "test@test.com"

    def test_extract_csv_empty_file(self, sample_dir):
        """CSV with only headers should return empty list."""
        path = os.path.join(sample_dir, "empty.csv")
        with open(path, "w", newline="") as f:
            f.write("name,dob,email\n")
        records = FormatDetector.extract(path)
        assert records == []

    def test_extract_csv_handles_empty_values(self, sample_dir):
        """CSV with empty cells should return empty strings."""
        path = os.path.join(sample_dir, "empty_vals.csv")
        with open(path, "w", newline="") as f:
            f.write("name,dob,email\n")
            f.write("Test,,test@test.com\n")
        records = FormatDetector.extract(path)
        assert records[0]["dob"] == ""


# ===========================================================================
# JSON Extraction Tests
# ===========================================================================


class TestJSONExtraction:
    """Test JSON file extraction."""

    def test_extract_json_list(self, json_file, sample_records):
        records = FormatDetector.extract(json_file)
        assert len(records) == len(sample_records)
        assert records[0]["full_name"] == "Muhammad Tayyab"

    def test_extract_json_single_object(self, json_single_record):
        """A JSON file with a single dict (not list) should return one record."""
        records = FormatDetector.extract(json_single_record)
        assert len(records) == 1
        assert records[0]["full_name"] == "Test User"

    def test_extract_json_nested_values(self, sample_dir):
        """JSON with nested values should preserve them."""
        path = os.path.join(sample_dir, "nested.json")
        data = [{"name": "Test", "address": {"city": "Lahore", "zip": "54000"}}]
        with open(path, "w") as f:
            json.dump(data, f)
        records = FormatDetector.extract(path)
        assert records[0]["address"] == {"city": "Lahore", "zip": "54000"}


# ===========================================================================
# XML Extraction Tests
# ===========================================================================


class TestXMLExtraction:
    """Test XML file extraction."""

    def test_extract_xml_basic(self, xml_file):
        records = FormatDetector.extract(xml_file)
        assert len(records) == 2
        assert records[0]["full_name"] == "Muhammad Tayyab"
        assert records[0]["email"] == "tayyab@example.com"
        assert records[1]["full_name"] == "Ali Ahmed"

    def test_extract_xml_flat_structure(self, sample_dir):
        """XML with flat structure (no child elements) should still work."""
        path = os.path.join(sample_dir, "flat.xml")
        with open(path, "w") as f:
            f.write('<?xml version="1.0"?><root><name>Test</name><value>123</value></root>')
        records = FormatDetector.extract(path)
        assert len(records) == 1
        assert records[0]["name"] == "Test"


# ===========================================================================
# TXT Extraction Tests
# ===========================================================================


class TestTXTExtraction:
    """Test TXT file extraction."""

    def test_extract_txt(self, txt_file):
        records = FormatDetector.extract(txt_file)
        assert len(records) == 1
        assert "content" in records[0]
        assert "plain text" in records[0]["content"]

    def test_extract_txt_empty_file(self, sample_dir):
        """Empty TXT file should return one record with empty content."""
        path = os.path.join(sample_dir, "empty.txt")
        with open(path, "w") as f:
            f.write("")
        records = FormatDetector.extract(path)
        assert len(records) == 1
        assert records[0]["content"] == ""


# ===========================================================================
# XLSX Extraction Tests
# ===========================================================================


class TestXLSXExtraction:
    """Test XLSX file extraction."""

    def test_extract_xlsx_basic(self, xlsx_file, sample_records):
        records = FormatDetector.extract(xlsx_file)
        assert len(records) == len(sample_records)
        assert records[0]["full_name"] == "Muhammad Tayyab"

    def test_extract_xlsx_generates_column_names(self, sample_dir):
        """XLSX with None header cells should generate col_N names."""
        from openpyxl import Workbook

        path = os.path.join(sample_dir, "no_headers.xlsx")
        wb = Workbook()
        ws = wb.active
        ws.append([None, "dob", None])
        ws.append(["Test", "1990-01-01", "extra"])
        wb.save(path)
        wb.close()
        records = FormatDetector.extract(path)
        assert len(records) == 1
        assert "col_0" in records[0]
        assert "col_2" in records[0]


# ===========================================================================
# DOCX Extraction Tests
# ===========================================================================


class TestDOCXExtraction:
    """Test DOCX file extraction."""

    def test_extract_docx_table(self, docx_file):
        records = FormatDetector.extract(docx_file)
        assert len(records) == 2
        assert records[0]["full_name"] == "Muhammad Tayyab"
        assert records[1]["email"] == "ali@test.org"

    def test_extract_docx_no_table_uses_paragraphs(self, sample_dir):
        """DOCX without tables should extract paragraph text."""
        from docx import Document

        path = os.path.join(sample_dir, "no_table.docx")
        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("Migration Data")
        doc.save(path)
        records = FormatDetector.extract(path)
        assert len(records) == 1
        assert "Hello World" in records[0]["content"]


# ===========================================================================
# Format Override Tests
# ===========================================================================


class TestFormatOverride:
    """Test that explicit format parameter overrides auto-detection."""

    def test_explicit_format_overrides_extension(self, sample_dir):
        """Passing file_format should override extension-based detection."""
        path = os.path.join(sample_dir, "data.csv")
        with open(path, "w") as f:
            f.write('[{"name": "Test"}]')
        # File is .csv but content is JSON — force JSON extraction
        records = FormatDetector.extract(path, file_format=FileFormat.JSON)
        assert len(records) == 1
        assert records[0]["name"] == "Test"

    def test_unsupported_format_raises_error(self, csv_file):
        """Passing an unsupported format should raise ValueError."""
        with pytest.raises(ValueError, match="Unsupported format"):
            FormatDetector.extract(csv_file, file_format="parquet")
