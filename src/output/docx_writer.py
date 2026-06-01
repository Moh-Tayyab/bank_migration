from docx import Document

from ..models import MigrationResult


class DOCXWriter:
    def write(self, result: MigrationResult, output_path: str):
        doc = Document()
        doc.add_heading("Migration Report", level=0)
        doc.add_paragraph(f"Success: {'Yes' if result.success else 'No'}")
        doc.add_paragraph(f"Total Records: {result.total_records}")
        doc.add_paragraph(f"Processed: {result.processed}")
        doc.add_paragraph(f"Failed: {result.failed}")
        if result.error:
            doc.add_paragraph(f"Error: {result.error}")
        if result.records:
            doc.add_paragraph()
            doc.add_heading("Migrated Records", level=1)
            fieldnames = list(result.records[0].keys())
            table = doc.add_table(rows=1, cols=len(fieldnames))
            table.style = "Table Grid"
            for i, name in enumerate(fieldnames):
                table.rows[0].cells[i].text = name
            for record in result.records:
                row = table.add_row().cells
                for i, name in enumerate(fieldnames):
                    row[i].text = str(record.get(name, ""))
        doc.add_paragraph()
        doc.add_heading("Audit Trail", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        header = table.rows[0].cells
        header[0].text = "Event"
        header[1].text = "Record ID"
        header[2].text = "Bank Pair"
        header[3].text = "Timestamp"
        for entry in result.audit_trail:
            row = table.add_row().cells
            row[0].text = entry.event.value
            row[1].text = entry.record_id
            row[2].text = entry.bank_pair
            row[3].text = str(entry.timestamp)
        doc.save(output_path)
