import csv
import json
import os
from typing import Any, Dict, List, Optional

from .models import FileFormat


class FormatDetector:
    MIME_MAP = {
        "text/csv": FileFormat.CSV,
        "application/json": FileFormat.JSON,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": FileFormat.DOCX,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": FileFormat.XLSX,
        "text/xml": FileFormat.XML,
        "application/xml": FileFormat.XML,
        "text/plain": FileFormat.TXT,
    }

    @staticmethod
    def detect_format(filepath: str) -> FileFormat:
        ext = os.path.splitext(filepath)[1].lower()
        ext_map = {
            ".csv": FileFormat.CSV,
            ".json": FileFormat.JSON,
            ".docx": FileFormat.DOCX,
            ".xlsx": FileFormat.XLSX,
            ".xml": FileFormat.XML,
            ".txt": FileFormat.TXT,
        }
        return ext_map.get(ext, FileFormat.TXT)

    @staticmethod
    def extract(filepath: str, file_format: Optional[FileFormat] = None) -> List[Dict[str, Any]]:
        fmt = file_format or FormatDetector.detect_format(filepath)
        extractors = {
            FileFormat.CSV: FormatDetector._extract_csv,
            FileFormat.JSON: FormatDetector._extract_json,
            FileFormat.DOCX: FormatDetector._extract_docx,
            FileFormat.XLSX: FormatDetector._extract_xlsx,
            FileFormat.XML: FormatDetector._extract_xml,
            FileFormat.TXT: FormatDetector._extract_txt,
        }
        extractor = extractors.get(fmt)
        if not extractor:
            raise ValueError(f"Unsupported format: {fmt}")
        return extractor(filepath)

    @staticmethod
    def _extract_csv(filepath: str) -> List[Dict[str, Any]]:
        records = []
        with open(filepath, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for row in reader:
                records.append(
                    {k.strip() if k else f"extra_{i}": v.strip() if v else "" for i, (k, v) in enumerate(row.items())}
                )
        return records

    @staticmethod
    def _extract_json(filepath: str) -> List[Dict[str, Any]]:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict):
            return [data]
        return data

    @staticmethod
    def _extract_docx(filepath: str) -> List[Dict[str, Any]]:
        from docx import Document

        doc = Document(filepath)
        records = []
        headers = []
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if j == 0:
                    headers = cells
                else:
                    if len(cells) == len(headers):
                        records.append(dict(zip(headers, cells)))
        if not records:
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            records.append({"content": text})
        return records

    @staticmethod
    def _extract_xlsx(filepath: str) -> List[Dict[str, Any]]:
        import openpyxl

        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        ws = wb.active
        records = []
        headers = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                headers = [str(cell).strip() if cell else f"col_{j}" for j, cell in enumerate(row)]
            else:
                record = {}
                for j, cell in enumerate(row):
                    if j < len(headers):
                        record[headers[j]] = str(cell) if cell is not None else ""
                if record:
                    records.append(record)
        wb.close()
        return records

    @staticmethod
    def _extract_xml(filepath: str) -> List[Dict[str, Any]]:
        import defusedxml.ElementTree as safe_ET

        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read().strip()
        if not content:
            raise ValueError(f"XML file is empty: {filepath}")

        tree = safe_ET.parse(filepath)
        root = tree.getroot()
        records = []

        # Check if this is a World-Check style XML with nested structure
        first_record = next((child for child in root), None)
        if first_record is not None and any(len(sub) > 0 for sub in first_record):
            # Use nested flattening for structures like World-Check
            for child in root:
                record = FormatDetector._flatten_xml_element(child, "")
                if record:
                    records.append(record)
        else:
            # Use simple extraction for flat XML structures
            for child in root:
                record = {}
                for sub in child:
                    record[sub.tag] = sub.text or ""
                if record:
                    records.append(record)

        if not records:
            record = {}
            for sub in root:
                record[sub.tag] = sub.text or ""
            if record:
                records.append(record)
        return records

    @staticmethod
    def _flatten_xml_element(elem: Any, prefix: str, max_depth: int = 5) -> Dict[str, Any]:
        """Recursively flatten XML elements, creating prefixed field names for nested structures.

        Args:
            elem: XML element to flatten
            prefix: Current field name prefix
            max_depth: Maximum depth to recurse (prevents infinite loops)

        Returns:
            Flattened dictionary with composite field names
        """
        if max_depth <= 0:
            return {}

        record = {}
        # Add element's own attributes
        for attr, value in elem.attrib.items():
            # Clean attribute names (remove namespace URIs)
            clean_attr = attr.split("}")[-1] if "}" in attr else attr
            attr_key = f"{prefix}_{clean_attr}" if prefix else clean_attr
            record[attr_key] = value

        # Process child elements
        child_counts = {}  # Track counts for multiple children with same tag
        for child in elem:
            child_counts[child.tag] = child_counts.get(child.tag, 0) + 1

        child_index = {}  # Track current index for each tag
        for child in elem:
            # Build field name with prefix
            new_prefix = f"{prefix}_{child.tag}" if prefix else child.tag

            # If multiple children with same tag, add index (starting from 1)
            if child_counts[child.tag] > 1:
                child_index[child.tag] = child_index.get(child.tag, 0) + 1
                new_prefix = f"{new_prefix}_{child_index[child.tag]}"

            # Check if child has children (nested structure)
            if len(child) > 0:
                # Recurse into nested elements
                nested = FormatDetector._flatten_xml_element(child, new_prefix, max_depth - 1)
                record.update(nested)
            else:
                # Leaf node - add child's attributes first
                for attr, value in child.attrib.items():
                    clean_attr = attr.split("}")[-1] if "}" in attr else attr
                    if clean_attr != "nil":  # Skip xsi:nil
                        attr_key = f"{new_prefix}_{clean_attr}"
                        record[attr_key] = value
                # Then add text content if present
                if child.text and child.text.strip():
                    record[new_prefix] = child.text.strip()
                # If no text and no attributes, still create the field
                elif not child.attrib:
                    record[new_prefix] = ""

        # If no content and no attributes, use element tag itself
        if not record and prefix:
            record[prefix] = ""

        return record

    @staticmethod
    def _extract_txt(filepath: str) -> List[Dict[str, Any]]:
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
        return [{"content": content}]
